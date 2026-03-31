"""
标准号提取器
============

从字符串中精确提取标准号，支持多种格式和特殊情况。
支持LLM辅助提取（通过服务端转发调用portal）。

使用方法：
    from extractor import StandardNumberExtractor

    extractor = StandardNumberExtractor()
    results = extractor.extract("参考GB/T 19001-2016和ISO 9001进行检测")
"""

import re
import json
import requests
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
import logging

logger = logging.getLogger(__name__)


@dataclass
class ExtractedStandard:
    """提取的标准号结果"""

    original: str  # 原始匹配文本
    normalized: str  # 标准化格式
    organization: str  # 标准组织
    number: str  # 标准编号
    year: Optional[str]  # 年份
    confidence: float  # 置信度 0-1
    method: str  # 提取方法: regex/llm
    has_chapter: bool = False  # 是否有章节号被排除


class StandardNumberExtractor:
    """标准号提取器"""

    PATTERNS = [
        r"GB(?:/T)?\s*\d{4,5}[-.]\d{4}",
        r"GBT\s*\d{4,5}[-.]\d{4}",
        r"ISO\s*\d{4,5}[-:]\d{4}",
        r"IEC\s*\d{4,5}[-:]\d{4}",
        r"ASTM\s*[A-Z]?\d{2,5}[-/]\d{2,4}",
        r"HG(?:/T)?\s*\d{4}[-.]\d{4}",
        r"QB(?:/T)?\s*\d{4}[-.]\d{4}",
        r"JB(?:/T)?\s*\d{4}[-.]\d{4}",
        r"DB\d{2}(?:/T)?\s*\d{4}[-.]\d{4}",
    ]

    def __init__(self, llm_config_path: str = None):
        """
        初始化提取器

        Args:
            llm_config_path: LLM配置文件路径
        """
        self.llm_config = self._load_llm_config(llm_config_path)
        self._compile_patterns()

    def _load_llm_config(self, config_path: str) -> Optional[dict]:
        """加载LLM配置"""
        if config_path is None:
            try:
                config_path = Path(__file__).parent.parent.parent / "model_config.json"
            except NameError:
                return None

        try:
            if Path(config_path).exists():
                with open(config_path, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception as e:
            logger.warning(f"加载LLM配置失败: {e}")

        return None

    def _compile_patterns(self):
        """编译正则表达式模式"""

        # 分隔符（半角全角混用）
        self.separators = r"[-—–:：]"

        # 年号模式（4位或2位）
        self.year_pattern = r"(?:19|20)?\d{2}"

        # 标准组织列表
        self.orgs = [
            # 中国标准
            r"GB\s*/?\s*[TJD]?",
            r"GBT",
            r"GB",
            # 行业标准
            r"HG\s*/?\s*T?",
            r"JB\s*/?\s*T?",
            r"SJ\s*/?\s*T?",
            r"YB\s*/?\s*T?",
            r"TB\s*/?\s*T?",
            r"DL\s*/?\s*T?",
            r"JG\s*/?\s*T?",
            r"QB\s*/?\s*T?",
            r"NY\s*/?\s*T?",
            r"SC\s*/?\s*T?",
            r"SN\s*/?\s*T?",
            r"WS\s*/?\s*T?",
            r"YY\s*/?\s*T?",
            r"JC\s*/?\s*T?",
            r"MT\s*/?\s*T?",
            r"SL\s*/?\s*T?",
            r"CJ\s*/?\s*T?",
            r"GA\s*/?\s*T?",
            r"LY\s*/?\s*T?",
            r"HY\s*/?\s*T?",
            r"HS\s*/?\s*T?",
            # 地方标准
            r"DB\d+\s*/?\s*T?",
            # 计量标准
            r"JJ[GFG]\s*",
            # 国际标准
            r"ISO\s*/?\s*IEC\s*",
            r"ISO\s*",
            r"IEC\s*",
            r"ASTM\s*[A-Z]?\s*",
            r"EN\s*",
            r"BS\s*",
            r"DIN\s*",
            r"JIS\s*",
            r"ANSI\s*/?\s*[A-Z]*\s*",
            r"IEEE\s*",
            r"API\s*",
            r"ASME\s*",
            # 其他
            r"CECS\s*",
            r"CNS\s*",
        ]

        org_pattern = "|".join(self.orgs)
        bracket_content = r"[（(][^)）]*[)）]?"
        generic_prefix = r"(?:[A-Z]{1,8}(?:/[A-Z]{1,8})*(?:/T)?|DB\d+(?:/T)?|JJ[FG])"

        # 模式1: 标准格式（高置信度）
        self.pattern_standard = re.compile(
            r"("
            + r"(?:"
            + org_pattern
            + r"|"
            + generic_prefix
            + r")"
            + r"\s*"
            + r"(?:"
            + bracket_content
            + r")?"
            + r"\s*"
            + r"[\d]+(?:\.\d+)?"
            + r"(?:\.\d+)?"
            + r"(?:"
            + bracket_content
            + r")?"
            + r"\s*"
            + r"(?:"
            + self.separators
            + r")"
            + r"\s*"
            + r"(?:"
            + self.year_pattern
            + r")"
            + r")",
            re.IGNORECASE,
        )

        # 模式2: 无年号格式（中置信度）
        self.pattern_no_year = re.compile(
            r"("
            + r"(?:"
            + org_pattern
            + r"|"
            + generic_prefix
            + r")"
            + r"\s*"
            + r"(?:"
            + bracket_content
            + r")?"
            + r"\s*"
            + r"[\d]+(?:\.\d+)?"
            + r"(?:\.\d+)?"
            + r")",
            re.IGNORECASE,
        )

        # 模式2.5: 仅数字+年份（弱匹配，用于缺少前缀的输入）
        self.pattern_bare_standard = re.compile(
            r"(?<!\d)(\d{3,6}\s*[-—–:：/]\s*(?:19|20)?\d{2})(?!\d)",
            re.IGNORECASE,
        )

        # 模式3: 中文标准名 + 年号
        self.pattern_chinese = re.compile(
            r"(?:依据|参考|按照|根据|执行|符合|详见|参见|引用|见)"
            r"\s*"
            r"《?"
            r"([\u4e00-\u9fa5]{2,15}"
            r"(?:标准|规范|规程|手册|指南|方法|细则|规定|条例|办法|文件))"
            r"》?"
            r"\s*"
            r"(?:" + self.separators + r")?"
            r"\s*"
            r"((?:19[89]\d|20[0-2]\d))"
            r"(?:\s*年)?"
            r"(?:\s*第?\d+号文?)?"
        )

        # 章节号模式（需要排除）
        self.pattern_chapter = re.compile(
            r"(?:\s*第\d+(?:\.\d+)*(?:章|节|条|款|项)|\s+\d+(?:\.\d+)*(?:章|节|条|款|项))\s*$"
        )

    def extract(self, text: str, use_llm: bool = False) -> List[ExtractedStandard]:
        """
        从文本中提取标准号

        Args:
            text: 输入文本
            use_llm: 是否使用LLM辅助提取（默认False）

        Returns:
            提取结果列表
        """
        results = []
        seen = set()

        # 第一层：标准格式匹配（高置信度）
        for match in self.pattern_standard.finditer(text):
            original = match.group(1).strip()
            original, has_chapter = self._remove_chapter(original)

            if original in seen:
                continue
            seen.add(original)

            normalized = self._normalize(original)
            org, number, year = self._parse_standard(original)

            results.append(
                ExtractedStandard(
                    original=original,
                    normalized=normalized,
                    organization=org,
                    number=number,
                    year=year,
                    confidence=0.95,
                    method="regex",
                    has_chapter=has_chapter,
                )
            )

        # 第二层：无年号格式匹配（中置信度）
        for match in self.pattern_no_year.finditer(text):
            original = match.group(1).strip()

            if original in seen or self._is_substring_of_existing(original, results):
                continue
            seen.add(original)

            normalized = self._normalize(original)
            org, number, year = self._parse_standard(original)

            results.append(
                ExtractedStandard(
                    original=original,
                    normalized=normalized,
                    organization=org,
                    number=number,
                    year=year,
                    confidence=0.8,
                    method="regex",
                    has_chapter=False,
                )
            )

        # 第2.5层：缺少前缀但带年份的标准号（低置信度）
        for match in self.pattern_bare_standard.finditer(text):
            original = match.group(1).strip()
            original, has_chapter = self._remove_chapter(original)

            if original in seen or self._is_substring_of_existing(original, results):
                continue
            seen.add(original)

            normalized = self._normalize(original)
            org, number, year = self._parse_standard(original)

            results.append(
                ExtractedStandard(
                    original=original,
                    normalized=normalized,
                    organization=org,
                    number=number,
                    year=year,
                    confidence=0.6,
                    method="regex",
                    has_chapter=has_chapter,
                )
            )

        # 第三层：中文标准名匹配
        for match in self.pattern_chinese.finditer(text):
            original = match.group(1).strip()
            year = match.group(2) if len(match.groups()) > 1 else None

            full_name = f"{original}{year}" if year else original

            if full_name in seen:
                continue
            seen.add(full_name)

            results.append(
                ExtractedStandard(
                    original=full_name,
                    normalized=full_name,
                    organization="中文标准",
                    number="",
                    year=year,
                    confidence=0.7,
                    method="regex",
                    has_chapter=False,
                )
            )

        # 第四层：LLM辅助提取（仅处理疑似片段）
        if use_llm and self.llm_config:
            suspected_segments = self._find_suspected_segments(text, results)

            if suspected_segments:
                llm_results = self._extract_with_llm(suspected_segments, results)
                for r in llm_results:
                    if r.normalized not in seen:
                        seen.add(r.normalized)
                        results.append(r)

        return results

    def _find_suspected_segments(
        self, text: str, existing_results: List[ExtractedStandard]
    ) -> List[str]:
        """找出疑似包含标准号的文本片段"""
        segments = []

        matched_texts = set(r.original for r in existing_results)

        delimiters = r"[。；;\n]"
        parts = re.split(delimiters, text)

        for part in parts:
            part = part.strip()
            if len(part) < 5 or len(part) > 200:
                continue

            contains_matched = False
            for matched in matched_texts:
                if matched in part:
                    contains_matched = True
                    break

            if contains_matched:
                continue

            has_year = bool(re.search(r"(?:19|20)\d{2}", part))
            has_number = bool(re.search(r"\d{2,}", part))
            has_standard_keyword = bool(
                re.search(r"(标准|规范|规程|依据|参考|按照)", part)
            )

            if has_year or (has_number and has_standard_keyword):
                segments.append(part)

        return segments[:5]

    def _remove_chapter(self, text: str) -> Tuple[str, bool]:
        """移除章节号"""
        match = self.pattern_chapter.search(text)
        if match:
            return text[: match.start()].strip(), True
        return text, False

    def _is_substring_of_existing(
        self, text: str, results: List[ExtractedStandard]
    ) -> bool:
        """检查是否是已有结果的子串"""
        for r in results:
            if text in r.original:
                return True
        return False

    def _normalize(self, text: str) -> str:
        """标准化标准号格式"""
        if not text:
            return text

        text = text.translate(
            str.maketrans(
                "ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ０１２３４５６７８９／－：",
                "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789/-:",
            )
        )

        text = re.sub(r"[—–:：]", "-", text)
        text = re.sub(r"\s+", " ", text).strip()
        text = text.upper()
        text = re.sub(r"GBT\s*", "GB/T ", text)
        text = re.sub(r"GB\s*/\s*T", "GB/T", text)

        return text

    def _parse_standard(self, text: str) -> Tuple[str, str, Optional[str]]:
        """解析标准号组成部分"""
        normalized = self._normalize(text)

        match = re.match(
            r"([A-Z]{1,8}(?:/[A-Z]{1,8})*(?:/T)?)\s*([\d]+(?:\.[\d]+)?)\s*(?:[-:]([\d]{2,4}))?",
            normalized,
        )

        if match:
            return match.group(1), match.group(2), match.group(3)

        match = re.match(r"([\d]+(?:\.[\d]+)?)\s*(?:[-:]([\d]{2,4}))?", normalized)
        if match:
            return "", match.group(1), match.group(2)

        return "", text, None

    def _extract_with_llm(
        self, suspected_segments: List[str], existing_results: List[ExtractedStandard]
    ) -> List[ExtractedStandard]:
        """使用LLM辅助提取标准号"""
        results = []

        if not suspected_segments:
            return results

        try:
            existing_standards = [r.original for r in existing_results]

            prompt = f"""请判断以下文本片段中是否包含标准号，如果有请提取。

文本片段：
{json.dumps(suspected_segments, ensure_ascii=False, indent=2)}

已提取的标准号（不需要重复）：
{json.dumps(existing_standards, ensure_ascii=False)}

提取规则：
1. 标准号格式如：GB/T 19001-2016, ISO 9001, 化妆品安全技术规范2015
2. 如果片段中没有标准号，返回空数组
3. 只返回JSON数组，格式：[{{"original": "原始文本", "normalized": "标准化格式", "confidence": 0.8}}]

返回JSON："""

            response = self._call_llm(prompt)

            if response:
                llm_standards = json.loads(response)

                for s in llm_standards:
                    if isinstance(s, dict) and s.get("original"):
                        results.append(
                            ExtractedStandard(
                                original=s.get("original", ""),
                                normalized=s.get("normalized", s.get("original", "")),
                                organization="",
                                number="",
                                year=None,
                                confidence=s.get("confidence", 0.6),
                                method="llm",
                            )
                        )

        except Exception as e:
            logger.error(f"LLM提取失败: {e}")

        return results

    def _call_llm(self, prompt: str) -> Optional[str]:
        """调用LLM API - 通过服务端转发到portal"""
        # 尝试获取hub_request（从app模块全局）
        try:
            from app import hub_request
        except ImportError:
            logger.warning("无法获取hub_request，LLM调用跳过")
            return None

        try:
            resp = hub_request("POST", "/api/llm/extract", {
                "prompt": prompt,
                "temperature": 0.1,
                "max_tokens": 2000,
            })

            if resp and resp.status_code == 200:
                data = resp.json()
                content = data.get("content", "")

                json_match = re.search(r"\[[\s\S]*\]", content)
                if json_match:
                    return json_match.group(0)

                return content

        except Exception as e:
            logger.error(f"LLM API调用失败 (通过服务端转发): {e}")

        return None

    def extract_from_file(
        self, file_path: str, enable_ocr: bool = False
    ) -> List[ExtractedStandard]:
        """从文件中提取标准号"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

            if not text:
                for encoding in ["gbk", "gb2312", "utf-16"]:
                    try:
                        with open(file_path, "r", encoding=encoding) as f:
                            text = f.read()
                        break
                    except UnicodeDecodeError:
                        continue
                    except Exception as e:
                        logger.debug(f"文件编码{encoding}读取失败: {e}")
                        continue

            return self.extract(text)

        except Exception as e:
            logger.error(f"文件读取失败: {e}")
            return []


def extract_standards(text: str, use_llm: bool = False) -> List[Dict]:
    """
    从文本中提取标准号（便捷函数）

    Args:
        text: 输入文本
        use_llm: 是否使用LLM辅助

    Returns:
        标准号列表（字典格式）
    """
    extractor = StandardNumberExtractor()
    results = extractor.extract(text, use_llm=use_llm)

    return [
        {
            "original": r.original,
            "normalized": r.normalized,
            "organization": r.organization,
            "number": r.number,
            "year": r.year,
            "confidence": r.confidence,
            "method": r.method,
        }
        for r in results
    ]


if __name__ == "__main__":
    test_cases = [
        "参考GB/T 19001-2016进行质量管理体系认证",
        "依据GB 5030-2002 4.3章节进行检测",
        "按照ISO/IEC 17025:2017标准执行",
        "符合ASTM D1234-20要求",
        "参照化妆品分析手册2002年3号文",
        "执行JJG 196-2006计量检定规程",
        "依据DB44/T 123-2020地方标准",
        "参考GBT19001—2016（等效ISO9001）",
        "按照农产品检验规范98进行检测",
        "依据HG/T 3934-2007和JB/T 10391-2008标准",
    ]

    extractor = StandardNumberExtractor()

    print("=" * 60)
    print("标准号提取测试")
    print("=" * 60)

    for text in test_cases:
        print(f"\n输入: {text}")
        results = extractor.extract(text, use_llm=False)

        for r in results:
            print(f"  提取: {r.original}")
            print(f"  标准化: {r.normalized}")
            print(f"  置信度: {r.confidence}")
            print(f"  方法: {r.method}")
            if r.has_chapter:
                print(f"  (已排除章节号)")


class StandardExtractor:
    """兼容客户端接口的提取器包装类"""

    PATTERNS = [
        r"GB(?:/T)?\s*\d{4,5}[-.]\d{4}",
        r"GBT\s*\d{4,5}[-.]\d{4}",
        r"ISO\s*\d{4,5}[-:]\d{4}",
        r"IEC\s*\d{4,5}[-:]\d{4}",
        r"ASTM\s*[A-Z]?\d{2,5}[-/]\d{2,4}",
        r"HG(?:/T)?\s*\d{4}[-.]\d{4}",
        r"QB(?:/T)?\s*\d{4}[-.]\d{4}",
        r"JB(?:/T)?\s*\d{4}[-.]\d{4}",
        r"DB\d{2}(?:/T)?\s*\d{4}[-.]\d{4}",
    ]

    def __init__(self, progress_callback=None):
        self.progress_callback = progress_callback
        self._inner = StandardNumberExtractor()

    def _report(self, current: int, total: int, message: str, details: dict = None):
        if self.progress_callback:
            self.progress_callback(current, total, message, details)

    def _extract_from_text(self, text: str) -> List[Dict]:
        """从文本提取标准号，返回兼容格式"""
        results = self._inner.extract(text, use_llm=False)
        return [
            {
                "standard": r.original,
                "confidence": r.confidence,
            }
            for r in results
        ]

    def extract_from_file(self, file_path: str, enable_ocr: bool = False) -> List[Dict]:
        """从文件提取标准号"""
        from pathlib import Path

        suffix = Path(file_path).suffix.lower()

        if suffix == ".pdf":
            try:
                import fitz

                doc = fitz.open(file_path)
                total_pages = len(doc)
                all_results = []

                for i, page in enumerate(doc):
                    text = page.get_text()
                    if text.strip():
                        page_results = self._extract_from_text(text)
                        for r in page_results:
                            r["page"] = i + 1
                        all_results.extend(page_results)

                    self._report(
                        int((i + 1) / total_pages * 100),
                        100,
                        f"处理第 {i + 1}/{total_pages} 页",
                        {"page": i + 1, "total": total_pages},
                    )

                doc.close()
                return all_results
            except ImportError:
                pass

        elif suffix == ".docx":
            try:
                from docx import Document

                doc = Document(file_path)
                text_parts = []
                for para in doc.paragraphs:
                    text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_parts.append(cell.text)

                text = "\n".join(text_parts)
                return self._extract_from_text(text)
            except ImportError:
                pass

        elif suffix == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return self._extract_from_text(text)

        elif suffix == ".csv":
            try:
                import csv

                all_results = []

                with open(
                    file_path, "r", encoding="utf-8", newline="", errors="ignore"
                ) as f:
                    reader = csv.reader(f)
                    for row in reader:
                        for cell in row:
                            if cell and isinstance(cell, str):
                                cell_results = self._extract_from_text(cell)
                                all_results.extend(cell_results)

                return all_results
            except Exception as e:
                logger.error(f"CSV文件处理失败: {e}")
                return []

        elif suffix in [".xlsx", ".xls"]:
            try:
                all_results = []

                if suffix == ".xlsx":
                    from openpyxl import load_workbook

                    wb = load_workbook(filename=file_path, read_only=True)
                    for sheet in wb.sheetnames:
                        ws = wb[sheet]
                        for row in ws.iter_rows(values_only=True):
                            for cell in row:
                                if cell and isinstance(cell, str):
                                    cell_results = self._extract_from_text(cell)
                                    all_results.extend(cell_results)
                    wb.close()
                else:  # .xls
                    import xlrd

                    workbook = xlrd.open_workbook(file_path)
                    for sheet_idx in range(workbook.nsheets):
                        sheet = workbook.sheet_by_index(sheet_idx)
                        for row_idx in range(sheet.nrows):
                            for col_idx in range(sheet.ncols):
                                cell = sheet.cell_value(row_idx, col_idx)
                                if cell and isinstance(cell, str):
                                    cell_results = self._extract_from_text(cell)
                                    all_results.extend(cell_results)

                return all_results
            except Exception as e:
                logger.error(f"Excel文件处理失败: {e}")
                return []

        return []
